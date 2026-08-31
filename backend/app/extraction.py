from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Optional

import fitz
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
CONTENT_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}


class UnsupportedFileTypeError(ValueError):
    pass


class TextExtractionError(ValueError):
    pass


def normalize_extension(filename: str, content_type: Optional[str] = None) -> str:
    extension = Path(filename or "").suffix.lower()
    if extension in SUPPORTED_EXTENSIONS:
        return extension
    if content_type in CONTENT_TYPES:
        return CONTENT_TYPES[content_type]
    raise UnsupportedFileTypeError("Only PDF, DOCX, and TXT resumes are supported.")


def extract_text(file_bytes: bytes, extension: str) -> str:
    try:
        if extension == ".pdf":
            return extract_pdf_text(file_bytes)
        if extension == ".docx":
            return extract_docx_text(file_bytes)
        if extension == ".txt":
            return extract_txt_text(file_bytes)
    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError("The file could not be parsed. It may be corrupt or invalid.") from exc
    raise UnsupportedFileTypeError("Only PDF, DOCX, and TXT resumes are supported.")


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as document:
            return "\n".join(page.get_text().strip() for page in document).strip()
    except Exception as exc:
        raise TextExtractionError("The PDF could not be parsed. It may be corrupt or invalid.") from exc


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        table_cells.append(text)
        return "\n".join(paragraphs + table_cells).strip()
    except Exception as exc:
        raise TextExtractionError("The DOCX file could not be parsed. It may be corrupt or invalid.") from exc


def extract_txt_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise TextExtractionError("The TXT file could not be decoded.")
